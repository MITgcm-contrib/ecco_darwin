#!/usr/bin/env python3
"""
Generate a Word document (docs/NS-RAD_running_instructions.docx) with step-by-step
instructions for running the NS-RAD model and rebuilding all figures.

Content is authored here so the document is reproducible; re-run to regenerate.

    python tools/make_run_instructions_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "NS-RAD_running_instructions.docx"

ACCENT = RGBColor(0x1C, 0x5B, 0x86)
MUTED = RGBColor(0x52, 0x51, 0x4E)


def shade(par, fill="F2F2F0"):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(doc, text, fill="F2F2F0"):
    """Shaded monospace block (one paragraph; newlines become line breaks)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    shade(p, fill)
    lines = text.strip("\n").split("\n")
    for i, ln in enumerate(lines):
        if i:
            r = p.add_run()
            r.add_break()
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def kv_table(doc, rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            para = cells[j].paragraphs[0]
            run = para.add_run(val)
            if j == 0:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
    return t


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def build():
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- title ----
    t = doc.add_heading("NS-RAD — Running Instructions", level=0)
    for r in t.runs:
        r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sr = sub.add_run("NS-RAD — the North Slope River-Aquatic-Delta Model: a C-GEM-based "
                     "1-D reactive-transport model of four Alaskan North Slope rivers "
                     "(Colville, Kuparuk, Sagavanirktok, Canning), plus a synthetic "
                     "verification fixture. This guide covers running the model and "
                     "rebuilding every figure/PDF/movie.")
    sr.italic = True
    sr.font.color.rgb = MUTED

    # ---- 1. quick start ----
    h(doc, "1. Quick start — one command", 1)
    doc.add_paragraph("From the project root (the folder containing code/, "
                      "tools/, runs/, docs/), run:")
    code(doc, "cd /Users/carrolld/Documents/research/FORTE/m_files/NS_RAD\ntools/build_all.sh")
    doc.add_paragraph("That single wrapper does the whole pipeline end to end:")
    bullet(doc, "runs C-GEM for all four rivers (the slow part, ~15 min, all four in parallel);",
           "Model — ")
    bullet(doc, "regenerates every figure PDF (diagnostics, validation, geometry, schematic, "
           "summary, river networks);", "Figures — ")
    bullet(doc, "regenerates the movies (per-river along-channel + Hovmöller animations).",
           "Movies — ")
    doc.add_paragraph("A single figure failing does not abort the rest — the wrapper prints a "
                      "PASS / FAIL / SKIP summary at the end and exits non-zero if anything failed.")

    # ---- 2. flags ----
    h(doc, "2. Options", 1)
    kv_table(doc, [
        ["tools/build_all.sh", "Run all four rivers, then all PDFs + movies."],
        ["tools/build_all.sh --figures-only",
         "Skip the model run; rebuild PDFs + movies from existing runs (~2–3 min)."],
        ["tools/build_all.sh --with-idealized",
         "Also run + verify + figure the idealized verification fixture (adds ~15 min)."],
        ["SERIAL=1 tools/build_all.sh", "Run the four rivers one at a time (easier to read logs)."],
    ], ["command", "what it does"])

    # ---- 3. prerequisites ----
    h(doc, "3. Prerequisites (one-time setup)", 1)
    doc.add_paragraph("The model uses the conda/miniforge Python environment (Python 3.13). "
                      "Required packages:")
    bullet(doc, "numpy, numba (load-bearing — the hot loops are JIT-compiled), netCDF4, matplotlib.")
    bullet(doc, "ffmpeg — for the .mp4 movies. Without it, generate GIFs instead "
           "(tools/make_movies.py --gif).")
    bullet(doc, "python-docx — only needed to regenerate this document.")
    doc.add_paragraph("Install anything missing with:")
    code(doc, "conda install numba netcdf4 matplotlib ffmpeg\n"
              "pip install python-docx")
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("numba compiles the kernels on the first run of a session (a one-time ~30–60 s "
              "cost); subsequent runs reuse the cache.")

    # ---- 4. outputs ----
    h(doc, "4. What it produces", 1)
    kv_table(doc, [
        ["runs/definitive/<site>/output.nc", "Model state per river (time × distance NetCDF)."],
        ["docs/ns_rad_diagnostics.pdf", "Modelled fields, all four rivers."],
        ["docs/ns_rad_validation.pdf", "Model vs observations (needs runs/regression_bnd)."],
        ["docs/ns_rad_geometry.pdf", "Channel geometry per river."],
        ["docs/ns_rad_model_schematic.pdf", "Land-to-ocean framework + data flow + call graph."],
        ["docs/ns_rad_model_summary.pdf", "Headline summary."],
        ["docs/ns_rad_river_networks.pdf", "River-network maps."],
        ["docs/movies/<site>_evolution.mp4", "Along-channel profile animation, per river."],
        ["docs/movies/<site>_hovmoller.mp4", "Distance × time animation with a day cursor."],
    ], ["output", "description"])

    # ---- 5. running pieces individually ----
    h(doc, "5. Running individual pieces", 1)
    h(doc, "Just the model", 2)
    code(doc, "tools/run_sites.sh                 # all four rivers -> runs/definitive/<site>/\n"
              "tools/run_sites.sh kuparuk         # one river\n"
              "tools/run_sites.sh colville kuparuk  # a subset\n"
              "SERIAL=1 tools/run_sites.sh        # one at a time")
    doc.add_paragraph("Each river runs as its own process in its own directory "
                      "(runs/definitive/<site>/). A quick smoke test (2 model days) without "
                      "editing config:")
    code(doc, "CGEM_MAXT_DAYS=2 CGEM_WARMUP_DAYS=1 tools/run_sites.sh")

    h(doc, "Just the figures (from existing runs)", 2)
    code(doc, "python tools/make_diagnostics_pdf.py\n"
              "python tools/make_validation_pdf.py     # needs runs/regression_bnd\n"
              "python tools/make_geometry_pdf.py\n"
              "python tools/make_schematic_pdf.py\n"
              "python tools/make_summary_figures.py\n"
              "python tools/make_river_maps.py\n"
              "python tools/make_movies.py             # all four; --gif for GIFs")

    # ---- 6. idealized fixture ----
    h(doc, "6. The idealized verification fixture", 1)
    doc.add_paragraph("A fifth, synthetic site — idealized — drives the whole model on a known, "
                      "analytic, time-varying input (including time-varying boundary conditions) "
                      "to catch regressions. It is a test fixture, not a river.")
    code(doc, "python tools/build_idealized_forcings.py   # (re)generate the analytic forcings\n"
              "python tools/verify_idealized.py --check-only  # structural checks, no model run\n"
              "python tools/verify_idealized.py            # + short 5-day run (~1–2 min)\n"
              "python tools/verify_idealized.py --full     # + 220-day seasonal run (~10–15 min)\n"
              "python tools/make_idealized_verification_pdf.py   # 4-page diagnostics PDF\n"
              "python tools/make_movies.py --run runs/idealized idealized")
    doc.add_paragraph("The harness exits non-zero if any physical-invariant check fails. Full "
                      "write-up: docs/idealized_verification.md.")

    # ---- 7. notes / troubleshooting ----
    h(doc, "7. Notes & troubleshooting", 1)
    bullet(doc, "the four rivers run in parallel in ~15 min; figures add a couple of minutes; "
           "movies are the slowest figure step.", "Runtime — ")
    bullet(doc, "the validation PDF needs runs/regression_bnd/ (an independent air-temperature-"
           "boundary run produced separately, not by run_sites.sh). build_all.sh SKIPs it if that "
           "directory is empty.", "Validation — ")
    bullet(doc, "\"ValueError: fp and xp are not of the same length\" means a wrong-length CSV was "
           "picked up. The real forcings live in forcing/ (365 values each).",
           "Forcing error — ")
    bullet(doc, "if main.py fails at import, numba is missing — conda install numba.",
           "numba — ")
    bullet(doc, "no ffmpeg? Use python tools/make_movies.py --gif to write animated GIFs instead "
           "of .mp4.", "Movies — ")
    bullet(doc, "runs/ is organized by run-type then river: definitive/ (operational), "
           "regression_bnd/ (validation independence check), idealized/ (fixture).",
           "Layout — ")

    p = doc.add_paragraph()
    p.add_run("See CLAUDE.md in the project root for the full architecture, and docs/ for "
              "provenance, performance and the ice-model plan.").italic = True

    doc.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    build()
